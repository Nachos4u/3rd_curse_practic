# -*- coding: utf-8 -*-
"""Помощник для формирования документов по ГОСТ (ЕСПД) средствами python-docx.

Параметры оформления:
  - формат A4, поля: левое 30 мм, правое 15 мм, верхнее/нижнее 20 мм;
  - основной шрифт Times New Roman 14 пт, межстрочный интервал 1.5;
  - абзацный отступ 1.25 см, выравнивание по ширине;
  - нумерация страниц снизу по центру.
"""
from docx import Document
from docx.shared import Pt, Mm, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


class GostDoc:
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
        self._setup_page(self.doc.sections[0])

    # -- базовая настройка --------------------------------------------------
    def _setup_styles(self):
        style = self.doc.styles["Normal"]
        style.font.name = "Times New Roman"
        style.font.size = Pt(14)
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), "Times New Roman")
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        pf.space_after = Pt(0)
        pf.space_before = Pt(0)

    def _setup_page(self, section):
        section.page_width = Mm(210)
        section.page_height = Mm(297)
        section.left_margin = Mm(30)
        section.right_margin = Mm(15)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)

    def _set_run_font(self, run, size=14, bold=False, italic=False):
        run.font.name = "Times New Roman"
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), "Times New Roman")

    # -- элементы содержания ------------------------------------------------
    def para(self, text="", align="justify", indent=True, bold=False, italic=False, size=14):
        p = self.doc.add_paragraph()
        amap = {"justify": WD_ALIGN_PARAGRAPH.JUSTIFY, "center": WD_ALIGN_PARAGRAPH.CENTER,
                "left": WD_ALIGN_PARAGRAPH.LEFT, "right": WD_ALIGN_PARAGRAPH.RIGHT}
        p.alignment = amap[align]
        if indent:
            p.paragraph_format.first_line_indent = Cm(1.25)
        if text:
            run = p.add_run(text)
            self._set_run_font(run, size, bold, italic)
        return p

    def h1(self, text, page_break=True):
        if page_break:
            self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(12)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        self._set_run_font(run, 16, bold=True)
        # для попадания в автособираемое оглавление
        p.style = self.doc.styles["Heading 1"]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def h2(self, text):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        self._set_run_font(run, 14, bold=True)
        p.style = self.doc.styles["Heading 2"]
        run.font.color.rgb = RGBColor(0, 0, 0)
        return p

    def bullet(self, text):
        p = self.doc.add_paragraph(style="List Bullet")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        self._set_run_font(run)
        return p

    def numbered(self, text):
        p = self.doc.add_paragraph(style="List Number")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = p.add_run(text)
        self._set_run_font(run)
        return p

    def image(self, path, width_mm=160, caption=None):
        self.doc.add_picture(path, width=Mm(width_mm))
        self.doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = self.para(caption, align="center", indent=False, italic=True, size=12)
            p.paragraph_format.space_after = Pt(6)

    def table(self, headers, rows, widths=None, caption=None):
        if caption:
            p = self.para(caption, align="left", indent=False, italic=True, size=12)
            p.paragraph_format.space_before = Pt(6)
        t = self.doc.add_table(rows=1, cols=len(headers))
        t.style = "Table Grid"
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for i, h in enumerate(headers):
            cell = t.rows[0].cells[i]
            cell.paragraphs[0].text = ""
            run = cell.paragraphs[0].add_run(h)
            self._set_run_font(run, 12, bold=True)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for row in rows:
            cells = t.add_row().cells
            for i, val in enumerate(row):
                cells[i].paragraphs[0].text = ""
                run = cells[i].paragraphs[0].add_run(str(val))
                self._set_run_font(run, 12)
        if widths:
            for row in t.rows:
                for i, w in enumerate(widths):
                    row.cells[i].width = Mm(w)
        return t

    # -- титул, оглавление, колонтитул -------------------------------------
    def title_page(self, org, work_type, theme, author, group, supervisor, city, year):
        for _ in range(2):
            self.para("", indent=False)
        self.para(org, align="center", indent=False, bold=True)
        for _ in range(6):
            self.para("", indent=False)
        self.para(work_type, align="center", indent=False, bold=True, size=18)
        self.para("на тему:", align="center", indent=False)
        self.para(f"«{theme}»", align="center", indent=False, bold=True, size=16)
        for _ in range(5):
            self.para("", indent=False)
        p = self.para("", indent=False, align="right")
        run = p.add_run(f"Выполнил(а): {author}")
        self._set_run_font(run)
        p2 = self.para("", indent=False, align="right")
        run2 = p2.add_run(f"Группа: {group}")
        self._set_run_font(run2)
        p3 = self.para("", indent=False, align="right")
        run3 = p3.add_run(f"Руководитель: {supervisor}")
        self._set_run_font(run3)
        # сдвигаем город/год в низ страницы
        for _ in range(6):
            self.para("", indent=False)
        self.para(f"{city}, {year}", align="center", indent=False)

    def toc(self):
        self.doc.add_page_break()
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run("СОДЕРЖАНИЕ")
        self._set_run_font(run, 16, bold=True)
        p.paragraph_format.space_after = Pt(12)

        par = self.doc.add_paragraph()
        run = par.add_run()
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = 'TOC \\o "1-2" \\h \\z \\u'
        fld_sep = OxmlElement("w:fldChar"); fld_sep.set(qn("w:fldCharType"), "separate")
        hint = OxmlElement("w:t"); hint.text = "Оглавление обновляется по F9 (Обновить поле)."
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_sep)
        run._r.append(hint); run._r.append(fld_end)

    def page_numbers(self):
        section = self.doc.sections[0]
        footer = section.footer
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        fld_begin = OxmlElement("w:fldChar"); fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar"); fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin); run._r.append(instr); run._r.append(fld_end)
        self._set_run_font(run, 12)

    def save(self, path):
        self.page_numbers()
        self.doc.save(path)
